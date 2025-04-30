import os
from typing import Any, Dict, List, Optional

from app.filebase.parsers.base_parser import BaseParser
from app.filebase.vector.file_chunk import FileChunk
from app.logger import get_logger

logger = get_logger(__name__)


class DocxParser(BaseParser):
    """
    Word文档解析器，处理.docx文件
    """

    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS: List[str] = ['.docx']

    def can_handle(self, file_path: str) -> bool:
        """
        检查是否可以处理该文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            bool: 是否可以处理
        """
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        解析Word文档
        
        Args:
            file_path: 文件路径
            metadata: 附加元数据
            
        Returns:
            Dict[str, Any]: 解析结果，包括元数据、内容和FileChunk对象列表
        """
        if metadata is None:
            metadata = {}

        # 添加解析器信息到元数据
        metadata['parser'] = 'docx'
        metadata['file_type'] = 'docx'

        try:
            import docx
            # 读取文档内容
            doc = docx.Document(file_path)

            # 提取文档属性
            core_properties = doc.core_properties
            if core_properties:
                if core_properties.title:
                    metadata['title'] = core_properties.title
                if core_properties.author:
                    metadata['author'] = core_properties.author
                if core_properties.created:
                    metadata['created'] = str(core_properties.created)
                if core_properties.modified:
                    metadata['modified'] = str(core_properties.modified)

            # 获取段落数并添加到元数据
            metadata['paragraph_count'] = len(doc.paragraphs)

            # 获取表格数并添加到元数据
            metadata['table_count'] = len(doc.tables)

            # 提取所有段落文本
            paragraphs_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # 忽略空段落
                    paragraphs_text.append(paragraph.text)

            # 提取表格内容
            tables_text = []
            for i, table in enumerate(doc.tables):
                table_content = []
                table_content.append(f"\n--- 表格 {i+1} ---")

                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_content.append(" | ".join(row_text))

                tables_text.append("\n".join(table_content))

            # 组合所有内容
            all_content = []
            if paragraphs_text:
                all_content.append("\n".join(paragraphs_text))
            if tables_text:
                all_content.append("\n".join(tables_text))

            content = "\n\n".join(all_content)

            # 按段落分组创建 FileChunk 对象
            file_chunks = []

            # 定义每个块的最大大小
            chunk_size = 3000  # 字符数
            current_chunk = []
            current_size = 0
            chunk_index = 0

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # 忽略空段落
                    # 如果当前段落加上当前块大小超过了最大块大小，创建一个新块
                    if current_size + len(paragraph.text) > chunk_size and current_chunk:
                        # 合并当前块内容
                        chunk_text = "\n".join(current_chunk)

                        # 创建 FileChunk
                        file_chunk = FileChunk(
                            text=chunk_text,
                            file_metadata=metadata,
                            chunk_metadata={
                                'paragraph_range': f"{chunk_index*chunk_size}-{chunk_index*chunk_size+len(current_chunk)}"
                            },
                            chunk_index=chunk_index,
                            total_chunks=0  # 临时值，最后更新
                        )
                        file_chunks.append(file_chunk)

                        # 重置块
                        current_chunk = []
                        current_size = 0
                        chunk_index += 1

                    # 添加当前段落到块
                    current_chunk.append(paragraph.text)
                    current_size += len(paragraph.text)

            # 处理最后一个块
            if current_chunk:
                chunk_text = "\n".join(current_chunk)

                # 创建 FileChunk
                file_chunk = FileChunk(
                    text=chunk_text,
                    file_metadata=metadata,
                    chunk_metadata={
                        'paragraph_range': f"{chunk_index*chunk_size}-{chunk_index*chunk_size+len(current_chunk)}"
                    },
                    chunk_index=chunk_index,
                    total_chunks=0  # 临时值，最后更新
                )
                file_chunks.append(file_chunk)

            # 如果有表格，将每个表格作为一个单独的块
            for i, table_text in enumerate(tables_text):
                file_chunk = FileChunk(
                    text=table_text,
                    file_metadata=metadata,
                    chunk_metadata={
                        'content_type': 'table',
                        'table_index': i + 1
                    },
                    chunk_index=len(file_chunks),
                    total_chunks=0  # 临时值，最后更新
                )
                file_chunks.append(file_chunk)

            # 更新所有块的 total_chunks
            total_chunks = len(file_chunks)
            for chunk in file_chunks:
                chunk.total_chunks = total_chunks

            return {
                'metadata': metadata,
                'content': content,
                'chunks': file_chunks
            }

        except ImportError:
            logger.error("未安装python-docx库，无法解析DOCX文件")
            return {
                'metadata': metadata,
                'content': "错误：解析DOCX文件需要安装python-docx库（pip install python-docx）",
                'chunks': []
            }
        except Exception as e:
            logger.error(f"解析DOCX文件 {file_path} 时出错: {e!s}")
            # 返回空内容
            return {
                'metadata': metadata,
                'content': "",
                'chunks': []
            } 
